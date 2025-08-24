#!/usr/bin/env python3
"""
Debug Process Script: последовательно выполняет все шаги из generate_report для отладки.
Перед запуском задайте свои пути к файлам и провайдера.
"""
import os
import sys
import logging
import traceback
import asyncio
import inspect
from difflib import get_close_matches
from docx import Document


from semantic_segmenter import segment_text_semantic

# --- Настройка логирования для отладки ---
logging.basicConfig(
    level=logging.DEBUG,  # DEBUG для максимальной детализации
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler("debug_run.log", encoding='utf-8')
    ]
)
logger = logging.getLogger("debug_run")

# --- Импорт модулей проекта ---
try:
    from utils.document_loader import DocumentLoader
    from langchain_utils.section_checker import SectionChecker
    from utils.parsers import (split_leaflet_sections,
                               split_ohlp_sections,
                               split_recommendations
                               )
    from utils.docx_writer import (
        build_replacements,
        replace_placeholders_in_doc,
        insert_recommendations,
        fill_comparison_table,
        save_with_timestamp,
    )
    logger.debug("Все необходимые модули успешно импортированы.")
except ImportError as e:
    logger.critical(f"Ошибка импорта модуля: {e}")
    logger.critical(
        "Убедитесь, что скрипт запускается из корневой директории проекта и все зависимости установлены."
    )
    sys.exit(1)

# ------------------- Настройки пользователя -------------------
TEST_PATH = "test_data/lv_gadobutrol.pdf"      # Замените на ваши реальные пути
REF_PATH = "test_data/lv_gadobutrol_etalon.pdf"
REC_PATH = "test_data/recommendations_lv.docx"
TEMPLATE_DIR = "templates"
TEMPLATE_NAME = "report_template_lv.docx"
PROVIDER = "yandex"
OUTPUT_DIR = "results"
PREFIX = "report_debug"
# --------------------------------------------------------------


def find_sections_table(doc: Document):
    """
    Безопасный поиск таблицы с разделами в шаблоне.
    """
    logger.debug("Поиск таблицы с разделами в шаблоне.")
    if len(doc.tables) > 2:
        tbl = doc.tables[2]
        if len(tbl.rows) > 1 and tbl.rows[1].cells:
            logger.debug("Найдена третья таблица, используем её.")
            return tbl
    logger.error("Таблица с разделами не найдена.")
    raise ValueError("Таблица с разделами не найдена в шаблоне.")


def run_debug_process():
    logger.info("--- Начало отладочного запуска ---")
    try:
        # 1) Извлечение текстов и метаданных
        logger.info("1) Извлечение текстов из документов…")
        loader_test = DocumentLoader(TEST_PATH)
        test_text = loader_test.load()
        logger.info(f"  → TEST ({TEST_PATH}): {len(test_text)} chars, "
                    f"type={loader_test.doc_type}, drug={loader_test.drug_name!r}")

        loader_ref = DocumentLoader(REF_PATH)
        ref_text = loader_ref.load()
        logger.info(f"  → REF  ({REF_PATH}):  {len(ref_text)} chars, "
                    f"type={loader_ref.doc_type}, drug={loader_ref.drug_name!r}")

        loader_rec = DocumentLoader(REC_PATH)
        rec_text = loader_rec.simple_load()
        logger.info(f"  → REC  ({REC_PATH}):  {len(rec_text)} chars")

        # 2) Инициализация SectionChecker
        logger.info(
            f"2) Инициализация SectionChecker с провайдером: {PROVIDER}")
        checker = SectionChecker(api_provider=PROVIDER)
        logger.info("  → Checker готов к работе")

        # 3) Загрузка шаблона и чтение секций
        logger.info("3) Загрузка шаблона и извлечение списка разделов…")
        tpl_path = os.path.join(TEMPLATE_DIR, TEMPLATE_NAME)
        if not os.path.exists(tpl_path):
            raise FileNotFoundError(f"Шаблон не найден: {tpl_path}")
        doc = Document(tpl_path)
        tbl = find_sections_table(doc)
        sections = [
            row.cells[0].text.strip()
            for row in tbl.rows[1:]
            if row.cells[0].text.strip()
        ]
        logger.info(f"  → {len(sections)} разделов: {sections}")
        

        # # 4) Разбиение на секции
        logger.info("4) Разбиение текстов по разделам…")
        if loader_test.doc_type == "leaflet":
           test_blocks = segment_text_semantic(test_text, sections)
           ref_blocks = segment_text_semantic(ref_text, sections)
        else:
            test_blocks = split_ohlp_sections(test_text, sections)
            ref_blocks = split_ohlp_sections(ref_text, sections)
        recs_blocks = split_recommendations(rec_text, sections)

        # 5) Проверка рекомендаций
        logger.info("5) Проверка рекомендаций для каждого раздела…")
        recommendations = {}
        for idx, sec in enumerate(sections, 1):
            logger.info(f"  {idx}/{len(sections)}: раздел '{sec}'…")
            try:
                actual = test_blocks.get(sec, "")
                match_keys = get_close_matches(
                    sec, recs_blocks.keys(), n=1, cutoff=0.7)
                if not match_keys:
                    logger.warning(
                        f" – рекомендаций не найдено, пропускаем")
                    continue
                rec_part = recs_blocks[match_keys[0]]
                if not rec_part.strip():
                    logger.debug(f" – рекомендации пусты, пропускаем")
                    continue
                result = checker.check_recommends(rec_part, actual)
                if inspect.isawaitable(result):
                    result = asyncio.run(result)
                logger.info(
                    f"    → получили рекомендацию (первые 100): {str(result)[:100]}...")
                recommendations[sec] = result
            except Exception as e:
                logger.error(f"Ошибка в разделе '{sec}': {e}")
                logger.debug(traceback.format_exc())
                recommendations[sec] = f"Ошибка: {e}"
        logger.info("  Все разделы обработаны.")

        # 6) Вставка рекомендаций и заполнение таблицы
        logger.info("6) Вставка рекомендаций и заполнение таблицы…")
        try:
            insert_recommendations(doc, recommendations)
            logger.debug("    – рекомендации вставлены")
        except Exception as e:
            logger.error(f"Ошибка вставки рекомендаций: {e}")
            logger.debug(traceback.format_exc())

        compare_data = [
            {
                "Раздел": sec,
                "Содержимое референтного документа": ref_blocks.get(sec, ""),
                "Содержимое тестируемого документа": test_blocks.get(sec, ""),
            }
            for sec in sections
        ]
        try:
            fill_comparison_table(doc, compare_data, table_index=2)
            logger.info("    – таблица обновлена")
        except Exception as e:
            logger.error(f"Ошибка заполнения таблицы: {e}")
            logger.debug(traceback.format_exc())

        # 7) Замена плейсхолдеров
        logger.info("7) Замена плейсхолдеров (имена, даты…)…")
        try:
            ref_name = loader_ref.drug_name
            test_name = loader_test.drug_name
            repls = build_replacements(ref_name, test_name)
            replace_placeholders_in_doc(doc, repls)
            logger.info(f"    – placeholders: {ref_name!r} → {test_name!r}")
        except Exception as e:
            logger.error(f"Ошибка замены плейсхолдеров: {e}")
            logger.debug(traceback.format_exc())

        # 8) Сохранение итогового отчёта
        logger.info("8) Сохранение итогового документа…")
        try:
            out_path = save_with_timestamp(
                doc, output_dir=OUTPUT_DIR, prefix=PREFIX)
            logger.info(f"    – отчёт сохранён: {out_path}")
        except Exception as e:
            logger.error(f"Ошибка сохранения: {e}")
            logger.debug(traceback.format_exc())

        logger.info("--- Отладочный запуск завершён успешно ---")

    except FileNotFoundError as fnf:
        logger.critical(f"Файл не найден: {fnf}")
    except Exception as e:
        logger.critical(f"Критическая ошибка: {e}")
        logger.debug(traceback.format_exc())
        sys.exit(1)


if __name__ == "__main__":
    run_debug_process()
