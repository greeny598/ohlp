import os
import re
from datetime import datetime
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from difflib import SequenceMatcher


def iter_paragraphs(element) -> List[Paragraph]:
    """
    Рекурсивно обходит все параграфы в переданном элементе:
    Document, _Cell (ячейка таблицы) или _Body (тело документа).
    """
    for para in element.paragraphs:
        yield para
    for table in getattr(element, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def insert_recommendations(doc: Document, recommendations: Dict[str, Any]) -> None:
    """
    Ищет в doc плейсхолдер {_RECOMMENDATIONS_}, очищает его и вставляет
    по одной строке для каждой записи recommendations со статусом != 'complied'.
    Каждая строка — новый абзац с:
      - отступом первой строки 1 см
      - выравниванием по ширине
      - ключ секции (жирным)
      - комментарий (с маленькой буквы, кавычки-ёлочки)
    """
    for paragraph in doc.paragraphs:
        if '{_RECOMMENDATIONS_}' in paragraph.text:
            paragraph.text = ''
            prev = paragraph

            for key, data in recommendations.items():
                print(key)
                if data.get('compliance') == 'complied':
                    continue
                comment = data.get('comments', '').strip()
                if not comment:
                    continue
                # lowercase first char
                comment = comment[0].lower() + comment[1:]
                # quotes to «»
                comment = re.sub(r'"([^\"]*)"', r'«\1»', comment)
                comment = re.sub(r"'([^']*)'", r'«\1»', comment)
                # build paragraph
                p_elm = OxmlElement('w:p')
                prev._p.addnext(p_elm)
                p = Paragraph(p_elm, prev._parent)
                fmt = p.paragraph_format
                fmt.first_line_indent = Cm(1)
                fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # section key bold
                run_section = p.add_run(f"{key}: ")
                run_section.bold = True
                run_section.font.size = Pt(14)
                # comment normal
                run_comment = p.add_run(comment)
                run_comment.font.size = Pt(14)
                prev = p
            break


def highlight_differences(a: str, b: str,
                          color_a: Tuple[int, int, int] = (0, 128, 0),
                          color_b: Tuple[int, int, int] = (255, 0, 0)) -> Tuple[List, List]:
    """
    Сравнивает строки a и b, возвращает два списка фрагментов:
    each element is (type, text, (color,))
    where type is 'plain' or 'highlight'.
    """
    matcher = SequenceMatcher(None, a, b)
    res_a, res_b = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        ta, tb = a[i1:i2], b[j1:j2]
        if tag == 'equal':
            res_a.append(('plain', ta))
            res_b.append(('plain', tb))
        else:
            if tag in ('replace', 'delete'):
                res_a.append(('highlight', ta, color_a))
            if tag in ('replace', 'insert'):
                res_b.append(('highlight', tb, color_b))
    return res_a, res_b


def replace_placeholders_in_doc(doc: Document, replacements: dict):
    """
    Заменяет все ключи из replacements на их значения в документе,
    во всех параграфах — включая таблицы любой вложенности.

    Логика форматирования:
      - Если замена произошла в ячейке таблицы: 12 pt, полужирный.
      - Если замена произошла вне таблиц: 14 pt.
    Подсветка спец-строк (<!-- formula-not-decoded -->, <!-- image -->) сохраняется.
    """
    def _in_table(para: Paragraph) -> bool:
        # параграф находится в ячейке таблицы, если его родитель — w:tc
        try:
            return para._p.getparent().tag.endswith('}tc')
        except Exception:
            return False

    # подготовим регулярку для поиска ключей
    pattern = None
    if replacements:
        # порядок ключей не важен, экранируем для безопасности
        pattern = re.compile("|".join(re.escape(k) for k in replacements.keys()))

    for para in iter_paragraphs(doc):
        orig_text = para.text

        # 1) ЗАМЕНА ПЛЕЙСХОЛДЕРОВ
        if pattern is not None and orig_text:
            replaced_text = pattern.sub(lambda m: replacements[m.group(0)], orig_text)

            if replaced_text != orig_text:
                in_table = _in_table(para)

                # Чистим все существующие run'ы и вставляем ОДИН новый run без дублирования
                for run in para.runs:
                    run.text = ""

                new_run = para.add_run(replaced_text)
                if in_table:
                    new_run.font.size = Pt(12)
                    new_run.bold = True
                else:
                    new_run.font.size = Pt(14)

        # 2) ПОДСВЕТКА спец-строк (выполняется и если замен не было)
        # Проходим по run'ам, разбиваем по таргетам, заново собираем с сохранением формата.
        targets = ["<!-- formula-not-decoded -->", "<!-- image -->"]
        if any(t in para.text for t in targets):
            pattern_highlight = re.compile("(" + "|".join(re.escape(t) for t in targets) + ")")
            new_fragments = []  # (text, formatting_dict, highlight_flag)

            for run in para.runs:
                text = run.text
                fmt = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                    'highlight': run.font.highlight_color
                }
                if text:
                    parts = pattern_highlight.split(text)
                    for part in parts:
                        if not part:
                            continue
                        if part in targets:
                            new_fragments.append((part, fmt, True))
                        else:
                            new_fragments.append((part, fmt, False))

            # очистка и пересборка
            for run in para.runs:
                run.text = ""

            for text, fmt, highlight in new_fragments:
                new_run = para.add_run(text)
                new_run.bold = fmt['bold']
                new_run.italic = fmt['italic']
                new_run.underline = fmt['underline']
                if fmt['font_name']:
                    new_run.font.name = fmt['font_name']
                if fmt['font_size']:
                    new_run.font.size = fmt['font_size']
                if highlight:
                    new_run.font.color.rgb = RGBColor(0, 0, 0)
                    new_run.font.highlight_color = WD_COLOR_INDEX.RED
                else:
                    if fmt['color']:
                        new_run.font.color.rgb = fmt['color']
                    if fmt['highlight']:
                        new_run.font.highlight_color = fmt['highlight']



def build_replacements(ref_name: str, test_name: str) -> Dict[str, str]:
    """
    Формирует словарь для replace_placeholders_in_doc:
      {{'_REF_NAME_': ref_name, '_TEST_NAME_': test_name, '_DATE_': today}}
    """
    today = datetime.now().strftime("%d.%m.%Y") + " г."
    return {
        "{_REF_NAME_}": ref_name,
        "{_TEST_NAME_}": test_name,
        "{_DATE_}": today,
    }


# --- вспомогательные функции для таблицы сравнения ---
def _clean_section_name(s: str) -> str:
    """
    Очищает строку: убирает ведущие и завершающие пробелы, звёздочки,
    кавычки, дефисы и прочие спецсимволы. Используется для названий разделов.
    """
    if not s:
        return ""
    # символы в начале
    s = re.sub(r'^[\s\*<>"«»№\.\-–—\u00A0]+', '', s.strip())
    # символы в конце
    s = re.sub(r'[\s\*<>"«»\.\-–—\u00A0]+$', '', s)
    return s.strip()


def _first_non_empty_line(text: str) -> str:
    """
    Возвращает первую непустую строку текста.
    """
    for ln in (text or '').splitlines():
        if ln.strip():
            return ln.strip()
    return ''


def _remove_first_line(text: str) -> str:
    """
    Возвращает текст без первой непустой строки (используется, чтобы
    убрать заголовок раздела из текста и не дублировать его в таблице).
    """
    if not text:
        return ""
    lines = text.splitlines()
    found = False
    new_lines: List[str] = []
    for ln in lines:
        if not found and ln.strip():
            found = True
            continue
        new_lines.append(ln)
    return "\n".join(new_lines).lstrip()


def _write_fragments(paragraph, frags: List[Tuple[str, str, Tuple[int, int, int]]], default_size_pt: int = 12) -> None:
    """
    Записывает список фрагментов (тип, текст, (r,g,b)) в параграф.
    typ может быть 'plain' или 'highlight'.
    Цвет применяется только для 'highlight'. Шрифт устанавливается стандартным
    размером и не выделяется полужирным.
    """
    # удалить существующие run'ы
    for run in paragraph.runs:
        run.text = ''
    for item in frags:
        if not item:
            continue
        kind = item[0]
        txt = item[1]
        color = item[2] if len(item) > 2 else None
        if not txt:
            continue
        run = paragraph.add_run(txt)
        run.font.size = Pt(default_size_pt)
        # не выделяем текст тела жирным
        run.bold = False
        if kind == 'highlight' and color:
            r, g, b = color
            run.font.color.rgb = RGBColor(r, g, b)

def fill_comparison_table(doc: Document,
                          recommended_sections: List[str],
                          ref_blocks: Dict[str, str],
                          test_blocks: Dict[str, str],
                          table_index: int = 2) -> None:
    """
    ЛИНЕЙНОЕ заполнение таблицы:
    i-й эталон → i-й референт → i-й тестируемый.
    Заголовки берутся ИЗ КЛЮЧЕЙ словарей ref_blocks/test_blocks.
    Текст разделов — из values.
    """
    table = doc.tables[table_index]

    # очищаем таблицу, оставляя только шапку (первую строку)
    while len(table.rows) > 1:
        tbl = table._tbl
        tbl.remove(tbl.tr_lst[-1])

    # превращаем dict → списки (ключ сохраняет порядок)
    ref_items = list(ref_blocks.items())
    test_items = list(test_blocks.items())

    total = len(recommended_sections)

    for i in range(total):

        # -----------------------------
        # 1. Заголовок + тело референтного блока
        # -----------------------------
        if i < len(ref_items):
            ref_title, ref_text = ref_items[i]
        else:
            ref_title, ref_text = "", ""

        # -----------------------------
        # 2. Заголовок + тело тестового блока
        # -----------------------------
        if i < len(test_items):
            test_title, test_text = test_items[i]
        else:
            test_title, test_text = "", ""

        # -----------------------------
        # 3. Заголовки для строки №2 (берём ИЗ КЛЮЧЕЙ, это важно!)
        # -----------------------------
        ref_header = _clean_section_name(ref_title)
        test_header = _clean_section_name(test_title)

        # -----------------------------
        # 4. Текст разделов (ничего не вырезаем)
        # -----------------------------
        ref_body = ref_text or ""
        test_body = test_text or ""

        # -----------------------------
        # 5. Подсветка различий
        # -----------------------------
        ref_frags, test_frags = highlight_differences(ref_body, test_body)

        # ==========================================================================
        # СТРОКА №1: ЭТАЛОННЫЙ ЗАГОЛОВОК
        # ==========================================================================
        row1 = table.add_row()
        merged = row1.cells[0]
        for c in row1.cells[1:]:
            merged = merged.merge(c)

        merged.text = ""
        p = merged.paragraphs[0]
        run = p.add_run(recommended_sections[i])
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ==========================================================================
        # СТРОКА №2: Заголовки из референтного и тестового документов
        # ==========================================================================
        row2 = table.add_row()

        # REF HEADER
        cell_ref = row2.cells[0]
        cell_ref.text = ""
        p_ref_h = cell_ref.paragraphs[0]
        run_ref_h = p_ref_h.add_run(ref_header)
        run_ref_h.bold = True
        run_ref_h.font.size = Pt(12)
        p_ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # TEST HEADER
        if len(row2.cells) > 1:
            cell_test = row2.cells[1]
            cell_test.text = ""
            p_test_h = cell_test.paragraphs[0]
            run_test_h = p_test_h.add_run(test_header)
            run_test_h.bold = True
            run_test_h.font.size = Pt(12)
            p_test_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ==========================================================================
        # СТРОКА №3: Сравнение текста (с подсветкой)
        # ==========================================================================
        row3 = table.add_row()

        # REF TEXT
        cell_ref_body = row3.cells[0]
        cell_ref_body.text = ""
        p_ref_body = cell_ref_body.paragraphs[0]
        _write_fragments(p_ref_body, ref_frags)

        # TEST TEXT
        if len(row3.cells) > 1:
            cell_test_body = row3.cells[1]
            cell_test_body.text = ""
            p_test_body = cell_test_body.paragraphs[0]
            _write_fragments(p_test_body, test_frags)



def save_with_timestamp(doc: Document,
                        filetype: str,
                        original_filename: str,
                        output_dir: str = "results",
                        prefix: str = "report") -> str:
    """
    Сохраняет doc в папке output_dir с названием
    prefix_DD_MM_YY(HH_MM_SS).docx, возвращает путь.
    """
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%d.%m.%y_(%H-%M-%S)")
    filename = f"{prefix}_{filetype}_{original_filename}_{ts}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path
