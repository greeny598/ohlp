import os
import re
from datetime import datetime
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt
from rapidfuzz import fuzz
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from difflib import SequenceMatcher


def iter_paragraphs(element) -> List[Paragraph]:
    """
    Рекурсивно обходит все параграфы в переданном элементе:
    Document, _Cell (ячейка таблицы) или _Body (тело документа).
    """
    for para in element.paragraphs:
        yield para
    for table in getattr(element, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def insert_recommendations(doc: Document, recommendations: Dict[str, Any]) -> None:
    """
    Ищет в doc плейсхолдер {_RECOMMENDATIONS_}, очищает его и вставляет
    по одной строке для каждой записи recommendations со статусом != 'complied'.
    Каждая строка — новый абзац с:
      - отступом первой строки 1 см
      - выравниванием по ширине
      - ключ секции (жирным)
      - комментарий (с маленькой буквы, кавычки-ёлочки)
    """
    for paragraph in doc.paragraphs:
        if '{_RECOMMENDATIONS_}' in paragraph.text:
            paragraph.text = ''
            prev = paragraph

            for key, data in recommendations.items():
                print(key)
                if data.get('compliance') == 'complied':
                    continue
                comment = data.get('comments', '').strip()
                if not comment:
                    continue
                # lowercase first char
                comment = comment[0].lower() + comment[1:]
                # quotes to «»
                comment = re.sub(r'"([^\"]*)"', r'«\1»', comment)
                comment = re.sub(r"'([^']*)'", r'«\1»', comment)
                # build paragraph
                p_elm = OxmlElement('w:p')
                prev._p.addnext(p_elm)
                p = Paragraph(p_elm, prev._parent)
                fmt = p.paragraph_format
                fmt.first_line_indent = Cm(1)
                fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # section key bold
                run_section = p.add_run(f"{key}: ")
                run_section.bold = True
                run_section.font.size = Pt(14)
                # comment normal
                run_comment = p.add_run(comment)
                run_comment.font.size = Pt(14)
                prev = p
            break


def highlight_differences(a: str, b: str,
                          color_a: Tuple[int, int, int] = (0, 128, 0),
                          color_b: Tuple[int, int, int] = (255, 0, 0)) -> Tuple[List, List]:
    """
    Сравнивает строки a и b, возвращает два списка фрагментов:
    each element is (type, text, (color,))
    where type is 'plain' or 'highlight'.
    """
    matcher = SequenceMatcher(None, a, b)
    res_a, res_b = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        ta, tb = a[i1:i2], b[j1:j2]
        if tag == 'equal':
            res_a.append(('plain', ta))
            res_b.append(('plain', tb))
        else:
            if tag in ('replace', 'delete'):
                res_a.append(('highlight', ta, color_a))
            if tag in ('replace', 'insert'):
                res_b.append(('highlight', tb, color_b))
    return res_a, res_b
       
def _norm_header_simple(s: str) -> str:
    if not s:
        return ""
    return (
        s.replace("\u00A0", " ")
         .replace("\n", " ")
         .replace("\r", " ")
         .strip()
    )

def set_cell_shading(cell, fill: str = "FF9999"):
    """
    Заливает ячейку таблицы цветом (по умолчанию светло-красным).
    """
    tc_pr = cell._tc.get_or_add_tcPr()

    # удалить старую заливку, если была
    for el in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(el)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)

    tc_pr.append(shd)


def replace_placeholders_in_doc(doc: Document, replacements: dict):
    """
    Заменяет все ключи из replacements на их значения в документе,
    во всех параграфах — включая таблицы любой вложенности.

    Логика форматирования:
      - Если замена произошла в ячейке таблицы: 12 pt, полужирный.
      - Если замена произошла вне таблиц: 14 pt.
    Подсветка спец-строк (<!-- formula-not-decoded -->, <!-- image -->) сохраняется.
    """
    def _in_table(para: Paragraph) -> bool:
        # параграф находится в ячейке таблицы, если его родитель — w:tc
        try:
            return para._p.getparent().tag.endswith('}tc')
        except Exception:
            return False

    # подготовим регулярку для поиска ключей
    pattern = None
    if replacements:
        # порядок ключей не важен, экранируем для безопасности
        pattern = re.compile("|".join(re.escape(k) for k in replacements.keys()))

    for para in iter_paragraphs(doc):
        orig_text = para.text

        # 1) ЗАМЕНА ПЛЕЙСХОЛДЕРОВ
        if pattern is not None and orig_text:
            replaced_text = pattern.sub(lambda m: replacements[m.group(0)], orig_text)

            if replaced_text != orig_text:
                in_table = _in_table(para)

                # Чистим все существующие run'ы и вставляем ОДИН новый run без дублирования
                for run in para.runs:
                    run.text = ""

                new_run = para.add_run(replaced_text)
                if in_table:
                    new_run.font.size = Pt(12)
                    new_run.bold = True
                else:
                    new_run.font.size = Pt(14)

        # 2) ПОДСВЕТКА спец-строк (выполняется и если замен не было)
        # Проходим по run'ам, разбиваем по таргетам, заново собираем с сохранением формата.
        targets = ["<!-- formula-not-decoded -->", "<!-- image -->"]
        if any(t in para.text for t in targets):
            pattern_highlight = re.compile("(" + "|".join(re.escape(t) for t in targets) + ")")
            new_fragments = []  # (text, formatting_dict, highlight_flag)

            for run in para.runs:
                text = run.text
                fmt = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                    'highlight': run.font.highlight_color
                }
                if text:
                    parts = pattern_highlight.split(text)
                    for part in parts:
                        if not part:
                            continue
                        if part in targets:
                            new_fragments.append((part, fmt, True))
                        else:
                            new_fragments.append((part, fmt, False))

            # очистка и пересборка
            for run in para.runs:
                run.text = ""

            for text, fmt, highlight in new_fragments:
                new_run = para.add_run(text)
                new_run.bold = fmt['bold']
                new_run.italic = fmt['italic']
                new_run.underline = fmt['underline']
                if fmt['font_name']:
                    new_run.font.name = fmt['font_name']
                if fmt['font_size']:
                    new_run.font.size = fmt['font_size']
                if highlight:
                    new_run.font.color.rgb = RGBColor(0, 0, 0)
                    new_run.font.highlight_color = WD_COLOR_INDEX.RED
                else:
                    if fmt['color']:
                        new_run.font.color.rgb = fmt['color']
                    if fmt['highlight']:
                        new_run.font.highlight_color = fmt['highlight']



def build_replacements(ref_name: str, test_name: str) -> Dict[str, str]:
    """
    Формирует словарь для replace_placeholders_in_doc:
      {{'_REF_NAME_': ref_name, '_TEST_NAME_': test_name, '_DATE_': today}}
    """
    today = datetime.now().strftime("%d.%m.%Y") + " г."
    return {
        "{_REF_NAME_}": ref_name,
        "{_TEST_NAME_}": test_name,
        "{_DATE_}": today,
    }


# --- вспомогательные функции для таблицы сравнения ---
def _clean_section_name(s: str) -> str:
    """
    Очищает строку: убирает ведущие и завершающие пробелы, звёздочки,
    кавычки, дефисы и прочие спецсимволы. Используется для названий разделов.
    """
    if not s:
        return ""
    # символы в начале
    s = re.sub(r'^[\s\*<>"«»№\.\-–—\u00A0]+', '', s.strip())
    # символы в конце
    s = re.sub(r'[\s\*<>"«»\.\-–—\u00A0]+$', '', s)
    return s.strip()

def _write_fragments(paragraph, frags: List[Tuple[str, str, Tuple[int, int, int]]], default_size_pt: int = 12) -> None:
    """
    Записывает список фрагментов (тип, текст, (r,g,b)) в параграф.
    typ может быть 'plain' или 'highlight'.
    Цвет применяется только для 'highlight'. Шрифт устанавливается стандартным
    размером и не выделяется полужирным.
    """
    # удалить существующие run'ы
    for run in paragraph.runs:
        run.text = ''
    for item in frags:
        if not item:
            continue
        kind = item[0]
        txt = item[1]
        color = item[2] if len(item) > 2 else None
        if not txt:
            continue
        run = paragraph.add_run(txt)
        run.font.size = Pt(default_size_pt)
        # не выделяем текст тела жирным
        run.bold = False
        if kind == 'highlight' and color:
            r, g, b = color
            run.font.color.rgb = RGBColor(r, g, b)
            
            



# =========================
# Markdown tables -> DOCX tables (шаг 1: простой парсер)
# =========================
_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")

def _split_md_blocks(text: str) -> List[Tuple[str, Any]]:
    """Разбивает текст на блоки ('text', str) и ('table', List[List[str]]).
    Поддерживает простые markdown-таблицы вида:
      | A | B |
      |---|---|
      | 1 | 2 |
    Всё, что между таблицами — один 'text' блок.
    """
    if not text:
        return []
    lines = text.splitlines()
    blocks: List[Tuple[str, Any]] = []
    buf: List[str] = []
    i = 0

    def flush_buf():
        nonlocal buf
        chunk = "\n".join(buf).strip()
        if chunk:
            blocks.append(("text", chunk))
        buf = []

    while i < len(lines):
        line = lines[i]
        # кандидат на заголовок таблицы: содержит '|'
        if "|" in line and i + 1 < len(lines) and _MD_TABLE_SEP_RE.match(lines[i + 1] or ""):
            # начинаем таблицу
            flush_buf()
            table_lines = [line, lines[i + 1]]
            i += 2
            # строки таблицы до первого пустого/непохожего
            while i < len(lines):
                ln = lines[i]
                if not ln.strip():
                    break
                if "|" not in ln:
                    break
                table_lines.append(ln)
                i += 1
            table = _parse_md_table_lines(table_lines)
            if table:
                blocks.append(("table", table))
            else:
                # если парсер не смог — возвращаем как текст
                blocks.append(("text", "\n".join(table_lines)))
            continue

        buf.append(line)
        i += 1

    flush_buf()
    return blocks


def _parse_md_table_lines(table_lines: List[str]) -> List[List[str]]:
    """Парсит markdown-таблицу из строк (header + sep + rows)."""
    if len(table_lines) < 2:
        return []
    header = table_lines[0]
    # table_lines[1] — разделитель
    body_lines = table_lines[2:] if len(table_lines) > 2 else []

    def split_row(row: str) -> List[str]:
        s = row.strip()
        # убрать крайние |
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        # простое разделение
        return [c.strip() for c in s.split("|")]

    rows: List[List[str]] = []
    rows.append(split_row(header))
    for bl in body_lines:
        rows.append(split_row(bl))

    # нормализация: одинаковое число колонок
    max_cols = max((len(r) for r in rows), default=0)
    if max_cols == 0:
        return []
    norm = []
    for r in rows:
        rr = r[:max_cols] + [""] * (max_cols - len(r))
        norm.append(rr)
    return norm


def _clear_cell(cell) -> None:
    """Аккуратно очищает ячейку от текста (таблицы python-docx в ячейке удалять сложно,
    но для нашего сценария мы создаём новую строку и пишем с нуля).
    """
    cell.text = ""
    # удалить лишние параграфы
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)



def _set_table_borders_dotted(table) -> None:
    """
    Делает границы таблицы явно видимыми пунктиром (dotted).
    Применяется ко всей таблице (tblBorders).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    # найти или создать элемент tblBorders
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    def _border(tag: str):
        el = tblBorders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tblBorders.append(el)
        el.set(qn("w:val"), "dotted")
        el.set(qn("w:sz"), "6")      # толщина
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    for t in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _border(t)


def _write_rich_with_tables(
    cell_ref,
    cell_test,
    ref_text: str,
    test_text: str,
    default_size_pt: int = 12,
) -> None:
    """Пишет в ячейки ref/test смешанный контент: текст + markdown-таблицы.
    Для текста — посимвольный diff (как раньше), для таблиц — diff по ячейкам, если структура совпадает.
    Фоллбек: если таблицы не совпадают по размеру — вставляем как plain text в таблицу без подсветки.
    """
    ref_blocks = _split_md_blocks(ref_text or "")
    test_blocks = _split_md_blocks(test_text or "")

    # выравнивание блоков по порядку (простое, без умной сопоставлялки)
    n = max(len(ref_blocks), len(test_blocks))

    _clear_cell(cell_ref)
    _clear_cell(cell_test)

    # начинаем писать в первый параграф
    pref = cell_ref.paragraphs[0]
    ptest = cell_test.paragraphs[0]

    def ensure_new_paragraph(cell, is_first: bool):
        if is_first:
            return cell.paragraphs[0]
        return cell.add_paragraph()

    def _maybe_parse_table_from_text(text_block: str):
        """Если текстовый блок на самом деле содержит markdown-таблицу,
        пытаемся распарсить **первую** таблицу и вернуть её как матрицу.

        Возвращает: (table_rows, remainder_text)
          - table_rows: List[List[str]] или []
          - remainder_text: текст без строк таблицы (если нашли)

        Это мягкий фоллбек, чтобы ситуация "слева таблица, справа пустая сетка"
        превращалась в "слева таблица, справа хотя бы текст" или даже таблица.
        """
        txt = (text_block or "").strip("\n")
        if not txt:
            return [], ""
        lines = txt.splitlines()
        # ищем начало markdown-таблицы (строка с | и следующая разделительная)
        for i in range(len(lines) - 1):
            if "|" not in lines[i]:
                continue
            if re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1]):
                # собираем до тех пор, пока строки похожи на таблицу
                chunk = [lines[i], lines[i + 1]]
                j = i + 2
                while j < len(lines) and ("|" in lines[j]):
                    chunk.append(lines[j])
                    j += 1
                try:
                    table_rows = _parse_md_table_lines(chunk)
                except Exception:
                    table_rows = []
                if table_rows:
                    remainder = (lines[:i] + lines[j:])
                    return table_rows, "\n".join(remainder).strip()
        return [], txt

    for idx in range(n):
        rblk = ref_blocks[idx] if idx < len(ref_blocks) else ("text", "")
        tblk = test_blocks[idx] if idx < len(test_blocks) else ("text", "")

        rtype, rdata = rblk
        ttype, tdata = tblk

        if rtype == "table" or ttype == "table":
            # Вставляем таблицы в каждую ячейку отдельно.
            # Если с одной стороны таблица, а с другой — текст, пытаемся
            # распарсить таблицу из текста (мягкий фоллбек).
            rtable = rdata if rtype == "table" else []
            ttable = tdata if ttype == "table" else []

            # если "таблица" не распознана с одной стороны, но там текст — попробуем извлечь
            r_remainder = ""
            t_remainder = ""
            if not rtable and rtype != "table":
                # rdata — текст
                rtable, r_remainder = _maybe_parse_table_from_text(rdata if isinstance(rdata, str) else "")
            if not ttable and ttype != "table":
                ttable, t_remainder = _maybe_parse_table_from_text(tdata if isinstance(tdata, str) else "")

            r_rows = len(rtable)
            r_cols = len(rtable[0]) if r_rows else 0
            t_rows = len(ttable)
            t_cols = len(ttable[0]) if t_rows else 0

            rows = max(r_rows, t_rows)
            cols = max(r_cols, t_cols)

            # если таблица всё ещё не распознана — пишем как текстовый блок с diff
            if rows == 0 or cols == 0:
                # fallback: пишем как текстовый блок с diff
                rtxt = (r_remainder or "")
                ttxt = (t_remainder or "")
                if rtable:
                    rtxt = ("\n".join(["\t".join(r) for r in (rtable or [])]) + "\n\n" + rtxt).strip()
                if ttable:
                    ttxt = ("\n".join(["\t".join(r) for r in (ttable or [])]) + "\n\n" + ttxt).strip()
                fr, ft = highlight_differences(rtxt, ttxt)
                pr = ensure_new_paragraph(cell_ref, is_first=(idx == 0))
                pt = ensure_new_paragraph(cell_test, is_first=(idx == 0))
                _write_fragments(pr, fr, default_size_pt=default_size_pt)
                _write_fragments(pt, ft, default_size_pt=default_size_pt)
                continue

            # добавляем небольшой отступ перед таблицей, если уже есть текст
            if idx > 0:
                cell_ref.add_paragraph("")
                cell_test.add_paragraph("")

            # ВАЖНО: если таблица есть только с одной стороны, не создаём "пустую сетку" на другой.
            # Вместо этого показываем там исходный текст (с подсветкой отличий относительно "плоского" представления таблицы).
            has_ref_tbl = rows > 0 and cols > 0 and r_rows > 0 and r_cols > 0
            has_test_tbl = rows > 0 and cols > 0 and t_rows > 0 and t_cols > 0

            ref_tbl = cell_ref.add_table(rows=rows, cols=cols) if has_ref_tbl else None
            test_tbl = cell_test.add_table(rows=rows, cols=cols) if has_test_tbl else None

            if ref_tbl is not None:
                _set_table_borders_dotted(ref_tbl)
            if test_tbl is not None:
                _set_table_borders_dotted(test_tbl)

            # ширины/стиль можно будет настроить позже; сейчас — безопасный шаг 1
            for ri in range(rows):
                for ci in range(cols):
                    rtxt = ""
                    ttxt = ""
                    if ri < r_rows and ci < r_cols:
                        rtxt = rtable[ri][ci]
                    if ri < t_rows and ci < t_cols:
                        ttxt = ttable[ri][ci]

                    # Всегда делаем посимвольный diff в ячейках.
                    fr, ft = highlight_differences(rtxt, ttxt)

                    if ref_tbl is not None:
                        pr = ref_tbl.cell(ri, ci).paragraphs[0]
                        pr.text = ""
                        _write_fragments(pr, fr, default_size_pt=default_size_pt)
                    if test_tbl is not None:
                        pt = test_tbl.cell(ri, ci).paragraphs[0]
                        pt.text = ""
                        _write_fragments(pt, ft, default_size_pt=default_size_pt)

            # если таблица была только с одной стороны — выводим контент второй стороны как текст (без пустой сетки)
            if ref_tbl is not None and test_tbl is None:
                # Плоское представление ref-таблицы vs test remainder/text
                flat_ref = "\n".join(["\t".join(r) for r in (rtable or [])]).strip()
                flat_test = (t_remainder or "").strip()
                fr, ft = highlight_differences(flat_ref, flat_test)
                pt = ensure_new_paragraph(cell_test, is_first=(idx == 0))
                pt.text = ""
                _write_fragments(pt, ft, default_size_pt=default_size_pt)
            elif test_tbl is not None and ref_tbl is None:
                flat_test = "\n".join(["\t".join(r) for r in (ttable or [])]).strip()
                flat_ref = (r_remainder or "").strip()
                fr, ft = highlight_differences(flat_ref, flat_test)
                pr = ensure_new_paragraph(cell_ref, is_first=(idx == 0))
                pr.text = ""
                _write_fragments(pr, fr, default_size_pt=default_size_pt)

            continue

        # оба текстовые — делаем посимвольный diff на уровне блока
        rtxt = rdata if isinstance(rdata, str) else ""
        ttxt = tdata if isinstance(tdata, str) else ""
        fr, ft = highlight_differences(rtxt, ttxt)

        pr = ensure_new_paragraph(cell_ref, is_first=(idx == 0))
        pt = ensure_new_paragraph(cell_test, is_first=(idx == 0))
        pr.text = ""
        pt.text = ""
        _write_fragments(pr, fr, default_size_pt=default_size_pt)
        _write_fragments(pt, ft, default_size_pt=default_size_pt)

def fill_comparison_table(doc: Document,
                          recommended_sections: List[str],
                          ref_blocks: Dict[str, str],
                          test_blocks: Dict[str, str],
                          table_index: int = 2) -> None:
    """
    Заполнение таблицы сравнения по эталонной структуре recommended_sections.

    1) Таблица строится строго по порядку разделов из recommended_sections.
    2) Блоки из ref_blocks и test_blocks сопоставляются с этими разделами
       по нечеткому совпадению заголовков (порог 0.8).
    3) В каждую тройку строк:
         - строка 1: эталонный заголовок (из recommended_sections);
         - строка 2: заголовки референтного и тестового документов;
         - строка 3: текст с подсветкой различий.
    4) Не сопоставленные разделы отправляются в блок «Дополнительные разделы»
       в конце таблицы.
    """

    table = doc.tables[table_index]

    # -----------------------------------------
    # 0. Очистка таблицы: оставляем только шапку
    # -----------------------------------------
    while len(table.rows) > 1:
        tbl = table._tbl
        tbl.remove(tbl.tr_lst[-1])

    # -----------------------------------------
    # 1. Вспомогательная функция сопоставления
    #    блоков документа с эталонными разделами
    # -----------------------------------------
    def match_blocks_to_sections(
        blocks: Dict[str, str],
        sections: List[str],
        threshold: float = 0.80,
    ) -> Tuple[Dict[str, Tuple[str, str]], List[Tuple[str, str]]]:
        """
        Возвращает:
          assigned: {canonical_section -> (header, text)}
          extras:   [ (header, text), ... ] — несопоставленные разделы
        """
        assigned: Dict[str, Tuple[str, str]] = {}
        extras: List[Tuple[str, str]] = []

        for raw_title, text in blocks.items():
            header = _clean_section_name(raw_title or "")
            if not header and not text:
                continue

            best_sec = None
            best_score = 0.0
            h_norm = header.lower()

            for sec in sections:
                score = SequenceMatcher(None, h_norm, sec.lower()).ratio()
                if score > best_score:
                    best_sec = sec
                    best_score = score

            if best_sec is not None and best_score >= threshold and best_sec not in assigned:
                assigned[best_sec] = (header, text or "")
            else:
                # либо не нашли хорошего совпадения, либо такой sec уже занят
                extras.append((header, text or ""))

        return assigned, extras

    # -----------------------------------------
    # 2. Сопоставляем ref/test блоки с эталонной структурой
    # -----------------------------------------
    ref_assigned, ref_extras = match_blocks_to_sections(ref_blocks, recommended_sections)
    test_assigned, test_extras = match_blocks_to_sections(test_blocks, recommended_sections)

    # -----------------------------------------
    # 3. Основные разделы по recommended_sections
    # -----------------------------------------
    for sec in recommended_sections:
        # Заголовок и текст референтного документа
        if sec in ref_assigned:
            ref_header, ref_body = ref_assigned[sec]
        else:
            ref_header, ref_body = "", ""

        # Заголовок и текст тестируемого документа
        if sec in test_assigned:
            test_header, test_body = test_assigned[sec]
        else:
            test_header, test_body = "", ""

        # Подсветка различий по ПОЛНОМУ тексту раздела
        ref_frags, test_frags = highlight_differences(ref_body, test_body)

        # ==========================================================================
        # СТРОКА №1: ЭТАЛОННЫЙ ЗАГОЛОВОК (из рекомендаций)
        # ==========================================================================
        row1 = table.add_row()
        merged = row1.cells[0]
        for c in row1.cells[1:]:
            merged = merged.merge(c)

        merged.text = ""
        p = merged.paragraphs[0]
        run = p.add_run(sec)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ==========================================================================
        # СТРОКА №2: Заголовки из референтного и тестового документов
        # ==========================================================================
        row2 = table.add_row()

        cell_ref = row2.cells[0]
        cell_ref.text = ""
        p_ref_h = cell_ref.paragraphs[0]

        if ref_header:
            run = p_ref_h.add_run(ref_header)
            run.bold = True
            run.font.size = Pt(12)

        # 🔴 если отличается от эталона или отсутствует
        if _norm_header_simple(ref_header) != _norm_header_simple(sec):
            set_cell_shading(cell_ref, "FF9999")

        p_ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER


        # TEST HEADER
        if len(row2.cells) > 1:
            cell_test = row2.cells[1]
            cell_test.text = ""
            p_test_h = cell_test.paragraphs[0]

            if test_header:
                run = p_test_h.add_run(test_header)
                run.bold = True
                run.font.size = Pt(12)

            # 🔴 если отличается от эталона или отсутствует
            if _norm_header_simple(test_header) != _norm_header_simple(sec):
                set_cell_shading(cell_test, "FF9999")

            p_test_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ==========================================================================
        # СТРОКА №3: Текст с подсветкой различий
        # ==========================================================================
        row3 = table.add_row()

        # REF/TST TEXT + TABLES (Шаг 1: markdown-таблицы)
        cell_ref_body = row3.cells[0]
        cell_test_body = row3.cells[1] if len(row3.cells) > 1 else None

        if cell_test_body is not None:
            _write_rich_with_tables(cell_ref_body, cell_test_body, ref_body, test_body)
        else:
            # fallback: если таблица сравнения вдруг одноколоночная
            cell_ref_body.text = ""
            p_ref_body = cell_ref_body.paragraphs[0]
            _write_fragments(p_ref_body, ref_frags)
# -----------------------------------------
    # 4. Дополнительные разделы (не нашедшие секцию в эталоне)
    # -----------------------------------------
    extras_total = max(len(ref_extras), len(test_extras))
    if extras_total > 0:
        # Заголовок блока "Дополнительные разделы"
        row_hdr = table.add_row()
        merged = row_hdr.cells[0]
        for c in row_hdr.cells[1:]:
            merged = merged.merge(c)
        merged.text = ""
        p_hdr = merged.paragraphs[0]
        run_hdr = p_hdr.add_run("РАЗДЕЛЫ, НЕ СООТВЕТСТВУЮЩИЕ РЕКОМЕНДАЦИЯМ")
        run_hdr.bold = True
        run_hdr.font.size = Pt(12)
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пары дополнительных блоков
        for i in range(extras_total):
            # Заголовки
            row_t = table.add_row()
            # Тела
            row_b = table.add_row()

            # REF extra
            if i < len(ref_extras):
                ref_h, ref_text = ref_extras[i]
            else:
                ref_h, ref_text = "", ""

            # TEST extra
            if i < len(test_extras):
                test_h, test_text = test_extras[i]
            else:
                test_h, test_text = "", ""

            # строка заголовков
            cell_ref_h = row_t.cells[0]
            cell_ref_h.text = ""
            p_ref_h = cell_ref_h.paragraphs[0]
            if ref_h:
                r = p_ref_h.add_run(ref_h)
                r.bold = True
                r.font.size = Pt(12)
            p_ref_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if len(row_t.cells) > 1:
                cell_test_h = row_t.cells[1]
                cell_test_h.text = ""
                p_test_h = cell_test_h.paragraphs[0]
                if test_h:
                    r2 = p_test_h.add_run(test_h)
                    r2.bold = True
                    r2.font.size = Pt(12)
                p_test_h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # строка текста (с подсветкой отличий между ref/test extra)
            cell_ref_b = row_b.cells[0]
            cell_ref_b.text = ""
            cell_test_b = row_b.cells[1]
            cell_test_b.text = ""

            p_ref_b = cell_ref_b.paragraphs[0]
            p_test_b = cell_test_b.paragraphs[0]

            ref_frags_extra, test_frags_extra = highlight_differences(ref_text, test_text)
            _write_fragments(p_ref_b, ref_frags_extra)
            _write_fragments(p_test_b, test_frags_extra)

def save_with_timestamp(doc: Document,
                        filetype: str,
                        original_filename: str,
                        output_dir: str = "results",
                        prefix: str = "report") -> str:
    """
    Сохраняет doc в папке output_dir с названием
    prefix_DD_MM_YY(HH_MM_SS).docx, возвращает путь.
    """
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%d.%m.%y_(%H-%M-%S)")
    filename = f"{prefix}_{filetype}_{original_filename}_{ts}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path
